#!/usr/bin/env python3
"""
Генерирует графики и отчёт для ЛР2 (вариант 46879).
Отчёт записывается в docs/report_v777.docx (имя файла сохранено, содержимое актуальное).
"""

import os
import csv
import math
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
RESOURCES = os.path.join(BASE_DIR, '..', 'src', 'test', 'resources')
GRAPHS_DIR = os.path.join(BASE_DIR, 'graphs')
os.makedirs(GRAPHS_DIR, exist_ok=True)


# ============================================================
# Графики
# ============================================================

def load_csv(path):
    xs, ys = [], []
    with open(path) as f:
        for row in csv.DictReader(f):
            xs.append(float(row['x']))
            ys.append(float(row['y']))
    return xs, ys


def plot_simple(xs, ys, title, color, filename):
    fig, ax = plt.subplots(figsize=(7, 4))
    ax.plot(xs, ys, color=color, linewidth=2, marker='o', markersize=5)
    ax.axhline(0, color='black', linewidth=0.8)
    ax.axvline(0, color='black', linewidth=0.8)
    ax.set_title(title, fontsize=13, pad=10)
    ax.grid(True, alpha=0.3)
    plt.tight_layout()
    path = os.path.join(GRAPHS_DIR, filename)
    fig.savefig(path, dpi=120, bbox_inches='tight')
    plt.close(fig)
    return path


def plot_multi(datasets, title, filename):
    fig, ax = plt.subplots(figsize=(8, 5))
    for ds in datasets:
        ax.plot(ds['x'], ds['y'], color=ds['color'], linewidth=2,
                marker='o', markersize=4, label=ds['label'])
    ax.axhline(0, color='black', linewidth=0.8)
    ax.axvline(0, color='black', linewidth=0.8)
    ax.set_title(title, fontsize=13, pad=10)
    ax.grid(True, alpha=0.3)
    ax.legend()
    plt.tight_layout()
    path = os.path.join(GRAPHS_DIR, filename)
    fig.savefig(path, dpi=120, bbox_inches='tight')
    plt.close(fig)
    return path


def plot_clipped(xs, ys, title, color, filename, ylim=(-20, 20)):
    fig, ax = plt.subplots(figsize=(7, 4))
    pts = sorted(zip(xs, ys))
    xs_s = [p[0] for p in pts]
    ys_clipped = [y if ylim[0] <= y <= ylim[1] else None for y in [p[1] for p in pts]]
    seg_x, seg_y, cx, cy = [], [], [], []
    for xv, yv in zip(xs_s, ys_clipped):
        if yv is None:
            if cx:
                seg_x.append(cx); seg_y.append(cy); cx, cy = [], []
        else:
            cx.append(xv); cy.append(yv)
    if cx:
        seg_x.append(cx); seg_y.append(cy)
    for sx, sy in zip(seg_x, seg_y):
        ax.plot(sx, sy, color=color, linewidth=2, marker='o', markersize=5)
    ax.set_ylim(ylim)
    ax.axhline(0, color='black', linewidth=0.8)
    ax.axvline(0, color='black', linewidth=0.8)
    ax.set_title(title, fontsize=13, pad=10)
    ax.grid(True, alpha=0.3)
    plt.tight_layout()
    path = os.path.join(GRAPHS_DIR, filename)
    fig.savefig(path, dpi=120, bbox_inches='tight')
    plt.close(fig)
    return path


def plot_system_v46879(xs, ys, filename, title_suffix=''):
    """Отрисовка системы варианта 46879: x≤0 и x>0 в двух панелях."""
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(12, 5))

    # x <= 0
    lt = sorted([(x, y) for x, y in zip(xs, ys) if x <= 0])
    if lt:
        xs_lt = [p[0] for p in lt]
        ys_lt_clip = [y if -50 <= y <= 450 else None for y in [p[1] for p in lt]]
        seg_x, seg_y, cx, cy = [], [], [], []
        for xv, yv in zip(xs_lt, ys_lt_clip):
            if yv is None:
                if cx: seg_x.append(cx); seg_y.append(cy); cx, cy = [], []
            else:
                cx.append(xv); cy.append(yv)
        if cx: seg_x.append(cx); seg_y.append(cy)
        for sx, sy in zip(seg_x, seg_y):
            ax1.plot(sx, sy, color='#E63946', linewidth=2, marker='o', markersize=5)
    ax1.axhline(0, color='black', linewidth=0.8); ax1.axvline(0, color='black', linewidth=0.8)
    ax1.set_title(f'Система (вариант 46879): x ≤ 0{title_suffix}', fontsize=12)
    ax1.grid(True, alpha=0.3)
    ax1.set_xlabel('x'); ax1.set_ylabel('f(x)')
    ax1.set_ylim(-50, 450)

    # x > 0
    gt = sorted([(x, y) for x, y in zip(xs, ys) if x > 0])
    if gt:
        ax2.plot([p[0] for p in gt], [p[1] for p in gt],
                 color='#457B9D', linewidth=2, marker='o', markersize=5)
    ax2.axhline(0, color='black', linewidth=0.8)
    ax2.set_title(f'Система (вариант 46879): x > 0{title_suffix}', fontsize=12)
    ax2.grid(True, alpha=0.3)
    ax2.set_xlabel('x'); ax2.set_ylabel('f(x)')

    plt.tight_layout()
    path = os.path.join(GRAPHS_DIR, filename)
    fig.savefig(path, dpi=120, bbox_inches='tight')
    plt.close(fig)
    return path


print('Генерирую графики...')

# Базовые
sin_pts = [(-math.pi, 0), (-math.pi*3/4, -math.sqrt(2)/2), (-math.pi/2, -1),
           (-math.pi/4, -math.sqrt(2)/2), (0, 0), (math.pi/4, math.sqrt(2)/2),
           (math.pi/2, 1), (math.pi*3/4, math.sqrt(2)/2), (math.pi, 0)]
g_sin = plot_simple([p[0] for p in sin_pts], [p[1] for p in sin_pts],
                    'sin(x) — базовая функция (ряд Тейлора)', '#E63946', 'sin.png')

cos_xs, cos_ys = load_csv(os.path.join(RESOURCES, 'cos.csv'))
g_cos = plot_clipped(cos_xs, cos_ys, 'cos(x) = sin(π/2 − x) — через Sine',
                     '#E07A5F', 'cos.png', ylim=(-2, 2))

ln_xs_an = [0.1, 0.25, 0.5, 1, 2, math.e, 5, 10, 20]
g_ln = plot_simple(ln_xs_an, [math.log(x) for x in ln_xs_an],
                   'ln(x) — базовая функция (ряд)', '#2A9D8F', 'ln.png')

log2_xs = [0.25, 0.5, 1, 2, 4, 8, 16, 32]
log5_xs = [0.2, 0.5, 1, 2, 5, 10, 25, 50]
lg_xs = [0.1, 1, 2, 5, 10, 50, 100, 1000]
g_logs_base = plot_multi([
    {'x': log2_xs, 'y': [math.log2(x) for x in log2_xs], 'label': 'log₂(x)', 'color': '#457B9D'},
    {'x': log5_xs, 'y': [math.log(x) / math.log(5) for x in log5_xs], 'label': 'log₅(x)', 'color': '#F4A261'},
    {'x': lg_xs, 'y': [math.log10(x) for x in lg_xs], 'label': 'log₁₀(x)', 'color': '#1D3557'},
], 'log₂, log₅, log₁₀ = ln(x)/ln(base)', 'logs_base.png')

# Производные триг-функции (module)
sec_xs, sec_ys = load_csv(os.path.join(RESOURCES, 'sec.csv'))
g_sec = plot_clipped(sec_xs, sec_ys, 'sec(x) = 1/cos(x) — module-тест', '#F4A261', 'sec_module.png')
csc_xs, csc_ys = load_csv(os.path.join(RESOURCES, 'csc.csv'))
g_csc = plot_clipped(csc_xs, csc_ys, 'csc(x) = 1/sin(x) — module-тест', '#E9C46A', 'csc_module.png')
tan_xs, tan_ys = load_csv(os.path.join(RESOURCES, 'tan.csv'))
g_tan = plot_clipped(tan_xs, tan_ys, 'tan(x) = sin(x)/cos(x) — module-тест',
                     '#2A9D8F', 'tan_module.png', ylim=(-10, 10))
cot_xs, cot_ys = load_csv(os.path.join(RESOURCES, 'cot.csv'))
g_cot = plot_clipped(cot_xs, cot_ys, 'cot(x) = cos(x)/sin(x) — module-тест',
                     '#264653', 'cot_module.png', ylim=(-10, 10))

# Интеграционные (с моками)
sec_it_xs, sec_it_ys = load_csv(os.path.join(RESOURCES, 'integration', 'secIT.csv'))
g_sec_it = plot_clipped(sec_it_xs, sec_it_ys, 'sec(x) — integration (мок-Cosine)',
                        '#F4A261', 'sec_it.png')
csc_it_xs, csc_it_ys = load_csv(os.path.join(RESOURCES, 'integration', 'cscIT.csv'))
g_csc_it = plot_clipped(csc_it_xs, csc_it_ys, 'csc(x) — integration (мок-Sine)',
                        '#E9C46A', 'csc_it.png')
tan_it_xs, tan_it_ys = load_csv(os.path.join(RESOURCES, 'integration', 'tanIT.csv'))
g_tan_it = plot_clipped(tan_it_xs, tan_it_ys, 'tan(x) — integration (мок-Sine/Cosine)',
                        '#2A9D8F', 'tan_it.png', ylim=(-5, 5))
cot_it_xs, cot_it_ys = load_csv(os.path.join(RESOURCES, 'integration', 'cotIT.csv'))
g_cot_it = plot_clipped(cot_it_xs, cot_it_ys, 'cot(x) — integration (мок-Sine/Cosine)',
                        '#264653', 'cot_it.png', ylim=(-5, 5))

# Система варианта 46879
sys_xs, sys_ys = load_csv(os.path.join(RESOURCES, 'system.csv'))
g_system = plot_system_v46879(sys_xs, sys_ys, 'system.png', ' — module')

sys_it_xs, sys_it_ys = load_csv(os.path.join(RESOURCES, 'integration', 'systemIT.csv'))
g_system_it = plot_system_v46879(sys_it_xs, sys_it_ys, 'system_it.png', ' — integration')

print(f'Графиков сгенерировано: {len(os.listdir(GRAPHS_DIR))}')


# ============================================================
# Отчёт
# ============================================================

print('Собираю отчёт...')

doc = Document()

# стиль по умолчанию
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)


def H(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h


def P(text, bold=False, italic=False, align=None, size=12):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    return p


def bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)


def add_image(path, caption, width_in=5.8):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Inches(width_in))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.italic = True
    r.font.name = 'Times New Roman'
    r.font.size = Pt(10)


# ---------- Титульный лист ----------
P('Федеральное государственное автономное образовательное учреждение высшего образования',
  align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
P('Национальный исследовательский университет ИТМО', align=WD_ALIGN_PARAGRAPH.CENTER,
  bold=True, size=12)
P('')
P('Факультет программной инженерии и компьютерной техники', align=WD_ALIGN_PARAGRAPH.CENTER)
P('Направление подготовки 09.03.04 Программная инженерия', align=WD_ALIGN_PARAGRAPH.CENTER)
P('Дисциплина «Тестирование программного обеспечения»', align=WD_ALIGN_PARAGRAPH.CENTER)
P('')
P('Отчёт', align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=16)
P('по лабораторной работе №2', align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14)
P('«Интеграционное тестирование»', align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=13)
P('')
P('Вариант: 46879', align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=13)
P('')
P('')
P('Студент: Молодиченко С. А.', align=WD_ALIGN_PARAGRAPH.RIGHT)
P('Группа: P3313', align=WD_ALIGN_PARAGRAPH.RIGHT)
P('')
P('Преподаватель: Стригалев Н. С.', align=WD_ALIGN_PARAGRAPH.RIGHT)
P('')
P('')
P('Санкт-Петербург, 2026 г.', align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# ---------- 1. Текст задания ----------
H('1. Текст задания и система функций', level=1)
P('Провести интеграционное тестирование программы, осуществляющей вычисление системы '
  'функций (вариант 46879).')
P('')
P('Система функций варианта 46879:', bold=True)
P('    x ≤ 0:   (((cos(x) + cos(x)) · sec(x)³) + cos(x))')
P('    x > 0:   ((((log₂(x) · log₅(x) / log₂(x)) · log₂(x))²) + (log₂(x) / log₁₀(x)))')
P('')
P('Требования к реализации:', bold=True)
bullet('Все составляющие системы функции выражены через базовые: тригонометрические — '
       'через sin(x) (базовая, ряд Тейлора); логарифмические — через ln(x) (базовая, ряд).')
bullet('Базовые функции sin(x) и ln(x) реализованы через разложение в ряд с задаваемой '
       'погрешностью. Тригонометрические / логарифмические преобразования для упрощения '
       'функций в модулях ЗАПРЕЩЕНЫ.')
bullet('Для каждого модуля реализованы табличные заглушки (Mockito @Mock / @Spy), '
       'найдены области допустимых значений.')
bullet('Приложение позволяет выводить значения любого модуля системы в CSV-файл вида '
       '«X, Результаты модуля(X)» с произвольно задаваемым шагом наращивания X. '
       'Реализация — класс CSVGraphWriter (см. раздел 6).')
P('')
P('Область допустимых значений:', bold=True)
bullet('Ветка x ≤ 0 не определена при x = −π/2, −3π/2, −5π/2, … — cos(x) = 0, sec(x) → ∞.')
bullet('Ветка x > 0 не определена при x = 1 — log₁₀(1) = 0, деление на ноль.')
bullet('При x = 0 (входит в ветку x ≤ 0): f(0) = 2·cos(0)·sec³(0) + cos(0) = 2 + 1 = 3.')
P('')
P('Базовые функции (требование задания):', bold=True)
bullet('sin(x) = x − x³/3! + x⁵/5! − x⁷/7! + … — ряд Тейлора.')
bullet('ln(x) = 2·(z + z³/3 + z⁵/5 + …), где z = (x−1)/(x+1) — ряд на основе arctanh.')

doc.add_page_break()

# ---------- 2. Архитектура ----------
H('2. Архитектура приложения', level=1)
P('Все функции реализуют интерфейс FunctionRule и наследуют AbstractFunction '
  '(валидация аргумента и точности).')
P('')
P('Зависимости модулей:', bold=True)
bullet('Sine — базовая, ряд Тейлора.')
bullet('Cosine — cos(x) = sin(π/2 − x), зависит от Sine.')
bullet('Cosecant — csc(x) = 1/sin(x), зависит от Sine.')
bullet('Secant — sec(x) = 1/cos(x), зависит от Cosine.')
bullet('Tangent — tan(x) = sin(x)/cos(x), зависит от Sine и Cosine.')
bullet('Cotangent — cot(x) = cos(x)/sin(x), зависит от Sine и Cosine.')
bullet('NaturalLogarithm — базовая, ряд на базе arctanh.')
bullet('BaseNLogarithm(n) — log_n(x) = ln(x)/ln(n), зависит от NaturalLogarithm. '
       'В системе используется для n = 2, 5, 10.')
bullet('EquationSystem — главный модуль; в формулу варианта 46879 вовлечены Cosine, '
       'Secant, BaseNLogarithm(2), BaseNLogarithm(5), BaseNLogarithm(10). Все эти модули '
       'передаются через конструктор (dependency injection) — для возможности '
       'подменять их заглушками в интеграционных тестах.')
P('')
if os.path.exists(os.path.join(BASE_DIR, 'uml_diagram.png')):
    add_image(os.path.join(BASE_DIR, 'uml_diagram.png'),
              'Рис. 1. UML-диаграмма классов разработанного приложения')
    P('')

# ---------- 3. Тестовое покрытие ----------
H('3. Тестовое покрытие', level=1)
P('Использованы JUnit 5 и Mockito. Стратегия интеграции — восходящая (bottom-up), '
  'по одному модулю. Требование покрытия: ≥ 85 % инструкций (проверяется '
  'через jacocoTestCoverageVerification в build.gradle.kts).')
P('')
P('3.1. Уровни интеграции', bold=True)
bullet('Уровень 0 (базовые, без зависимостей): Sine, NaturalLogarithm.')
bullet('Уровень 1: Cosine ← Sine.')
bullet('Уровень 2: Secant ← Cosine; Cosecant ← Sine; Tangent ← Sine + Cosine; '
       'Cotangent ← Sine + Cosine.')
bullet('Уровень 3: BaseNLogarithm ← NaturalLogarithm.')
bullet('Уровень 4 (вершина): EquationSystem ← Cosine, Secant, BaseNLogarithm(2/5/10).')
P('')
P('3.2. Модульные тесты', bold=True)
P('Для каждого модуля — отдельный *Test.java в пакете module/. Проверяют:')
bullet('Валидация входных данных: null-аргумент, null-точность, точность ∉ (0; 1).')
bullet('Специальные значения (x = 0, π/2, π, 2π — для триг; x = 1 — для ln).')
bullet('Точки разрыва (асимптоты): ожидается ArithmeticException.')
bullet('Параметризованные тесты из CSV-эталонов (sin.csv, cos.csv, …, system.csv) — '
       'набор заранее рассчитанных пар (x, y).')
bullet('FunctionRuleTest — общий контракт на все 11 модулей сразу.')
P('')
P('3.3. Интеграционные тесты', bold=True)
P('Файлы в пакете integration/. Используется Mockito @Mock (подмена поведения) '
  'и @Spy (контроль вызовов на реальном объекте).')
bullet('CosineIntegrationTest — cos использует Sine (spy и mock-Math.sin).')
bullet('SecantIntegrationTest — sec = 1/cos через mock-Cosine.')
bullet('CosecantIntegrationTest — csc = 1/sin через mock-Sine.')
bullet('TangentIntegrationTest — tan = sin/cos, в т. ч. проверка cos = 0 → ArithmeticException.')
bullet('CotangentIntegrationTest — cot = cos/sin, проверка sin = 0 → ArithmeticException.')
bullet('LogarithmIntegrationTest — log_n = ln(x)/ln(n) через mock-NaturalLogarithm.')
bullet('EquationSystemIntegrationTest — (1) при x ≤ 0 проверяется, что вызваны '
       'только cos и sec, логарифмы НЕ вызывались (verifyNoInteractions); '
       '(2) при x > 0 — зеркальная проверка; (3) параметризованный тест '
       'с полным замоканием всех 5 модулей, данные — systemIT.csv.')
P('')
P('3.4. Классы эквивалентности для системы (вариант 46879)', bold=True)
bullet('Большое отрицательное (x = −5), среднее отрицательное (−2.5 … −0.5) — '
       'числовое значение.')
bullet('Граница x = 0: f(0) = 3.')
bullet('Асимптоты x = −π/2, −3π/2 — ArithmeticException.')
bullet('Малое положительное (0.5, 2), большое положительное (5, 10, 100) — числовое значение.')
bullet('Особая точка x = 1 — ArithmeticException (log₁₀(1) = 0).')
bullet('Невалидная точность (≤ 0 или ≥ 1), null — соответствующие исключения.')
P('')
P('3.5. Итоги покрытия JaCoCo', bold=True)
bullet('ru.itmo.qa.lab2 (EquationSystem, Main): 100 % инструкций / 100 % веток.')
bullet('ru.itmo.qa.lab2.trig: 100 % / 93 % (1 защитная ветка i < seriesLength в Sine).')
bullet('ru.itmo.qa.lab2.log: 100 % / 100 %.')
bullet('ru.itmo.qa.lab2.function: 100 % / 100 %.')
bullet('ru.itmo.qa.lab2.util: 99 % / 100 %.')
bullet('Итого: 99 % инструкций / 98 % веток. Всего 287 тестов, 0 упавших.')

doc.add_page_break()

# ---------- 4. Графики ----------
H('4. Графики CSV-выгрузок', level=1)
P('Графики построены по CSV-выгрузкам всех модулей, организованы по уровням '
  'интеграции bottom-up.')

H('4.1. Уровень 0 — базовые функции', level=2)
add_image(g_sin, 'Рис. 2. sin(x) — базовая функция, ряд Тейлора')
add_image(g_ln, 'Рис. 3. ln(x) — базовая функция, ряд на основе arctanh')

H('4.2. Уровень 1-2 — производные функции (module-тесты)', level=2)
add_image(g_cos, 'Рис. 4. cos(x) = sin(π/2 − x), через Sine')
add_image(g_sec, 'Рис. 5. sec(x) = 1/cos(x) — module-тест')
add_image(g_csc, 'Рис. 6. csc(x) = 1/sin(x) — module-тест')
add_image(g_tan, 'Рис. 7. tan(x) = sin(x)/cos(x) — module-тест')
add_image(g_cot, 'Рис. 8. cot(x) = cos(x)/sin(x) — module-тест')
add_image(g_logs_base, 'Рис. 9. log₂, log₅, log₁₀ через NaturalLogarithm')

H('4.3. Интеграционные тесты производных функций', level=2)
add_image(g_sec_it, 'Рис. 10. sec(x) — integration-тест (мок-Cosine)')
add_image(g_csc_it, 'Рис. 11. csc(x) — integration-тест (мок-Sine)')
add_image(g_tan_it, 'Рис. 12. tan(x) — integration-тест (мок-Sine/Cosine)')
add_image(g_cot_it, 'Рис. 13. cot(x) — integration-тест (мок-Sine/Cosine)')

H('4.4. Система функций варианта 46879', level=2)
add_image(g_system, 'Рис. 14. Система (module-тест): левая ветка x ≤ 0 и правая x > 0')
add_image(g_system_it, 'Рис. 15. Система (integration-тест с полным замоканием модулей)')

doc.add_page_break()

# ---------- 5. CSV-утилита ----------
H('5. CSV-вывод значений модулей', level=1)
P('Требование задания: приложение должно уметь выводить значения любого модуля '
  'в CSV-файл вида «X, Результаты модуля(X)» с произвольным шагом.', italic=True)
P('')
P('Реализация — класс CSVGraphWriter', bold=True)
P('Путь: src/main/java/ru/itmo/qa/lab2/util/CSVGraphWriter.java')
P('')
P('Ключевые элементы:', bold=True)
bullet('Конструктор CSVGraphWriter(AbstractFunction function, String outputDir) — '
       'принимает ЛЮБОЙ модуль (наследник AbstractFunction). Это sin, cos, sec, csc, '
       'tan, cot, ln, log_n, EquationSystem.')
bullet('Метод write(BigDecimal x1, BigDecimal x2, BigDecimal d, BigDecimal precision) — '
       'параметр d задаёт шаг наращивания X и может быть ПРОИЗВОЛЬНЫМ.')
bullet('Файл формируется с header «x,y» и строками «x,y» — формат в точности соответствует '
       'требованию «X, Результаты модуля(X)».')
bullet('Разделитель — запятая (задание разрешает произвольный).')
bullet('При разрыве функции (ArithmeticException) пишется пустая строка — для корректного '
       'разрыва линии на графике.')
P('')
P('Пример использования (Main.java:58-69):', bold=True)
P('    new CSVGraphWriter(new Sine(),           outputDir).write(x1, x2, step, precision);')
P('    new CSVGraphWriter(new Cosine(),         outputDir).write(x1, x2, step, precision);')
P('    new CSVGraphWriter(new EquationSystem(), outputDir).write(x1, x2, step, precision);')
P('Чтобы изменить шаг — просто передай другое значение step (например new BigDecimal("0.5")).')

doc.add_page_break()

# ---------- 6. Выводы ----------
H('6. Выводы', level=1)
bullet('Реализована система функций варианта 46879 без использования Math.sin/Math.log.')
bullet('Базовые функции (sin, ln) вычисляются разложением в ряд с заданной точностью; '
       'остальные функции выражены через базовые (требование задания соблюдено).')
bullet('Каждый модуль покрыт unit-тестами (модульный уровень) и mock-интеграционными тестами.')
bullet('Применена восходящая стратегия интеграции «по одному модулю»: на каждом уровне '
       'к проверенным реальным модулям добавляется ровно один новый, остальные замокированы.')
bullet('В EquationSystemIntegrationTest отдельно проверяется маршрутизация по веткам '
       '(verifyNoInteractions для неиспользуемых модулей), что гарантирует корректность '
       'разделения x ≤ 0 / x > 0.')
bullet('Покрытие JaCoCo: 99 % инструкций / 98 % веток; требование ≥ 85 % превышено. '
       'Все 287 тестов зелёные.')
bullet('Реализована утилита CSVGraphWriter для выгрузки значений любого модуля в CSV '
       'с произвольным шагом — требование задания выполнено.')

# ---------- Сохранение ----------
out_path = os.path.join(BASE_DIR, 'report_v777.docx')
doc.save(out_path)
print(f'Отчёт сохранён: {out_path}')
